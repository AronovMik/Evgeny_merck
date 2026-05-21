#!/usr/bin/env python3
"""Generate ShkolaIBS clean invitation (Section 1 only, for sending to doctors)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = "/home/user/Evgeny_merck/tasks/letters/2026-05-20_shkola-ibs_invitation.docx"

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)

# Поля страницы
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_para(text, bold=False, italic=False, size=11, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')


# ===== TITLE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("ШКОЛА ИБС")
trun.bold = True
trun.font.size = Pt(22)
trun.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Современные подходы к диагностике и терапии")
sr.italic = True
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mrun = meta.add_run("17 июня 2026  •  15:00–19:00  •  {Город}, {Площадка}")
mrun.font.size = Pt(11)
mrun.bold = True
mrun.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

doc.add_paragraph()

# ===== Body =====
add_para("Уважаемые коллеги!", bold=True, size=12)

add_para("17 июня 2026 года в {Город} ({Площадка}) с 15:00 до 19:00 пройдёт Школа ИБС — "
         "образовательная встреча для кардиологов и терапевтов. Программа охватывает современную "
         "диагностику ишемической болезни сердца (включая хронические коронарные синдромы по ESC 2024 — "
         "от оценки симптомов и предтестовой вероятности до коронарного кальция, стресс-визуализации "
         "и решения о реваскуляризации) и обоснованный выбор терапии при стабильной ИБС, в том числе "
         "при наличии сопутствующих состояний. Рамка: КР МЗ РФ «Стабильная ИБС» 2024, "
         "КР МЗ РФ «АГ у взрослых» 2024, ESC 2024 CCS.")

add_para("Эксперты — два ведущих специалиста в совместном разборе:", bold=True)
add_bullet("Гражданкин Игорь Олегович — к.м.н., врач высшей категории, заведующий "
           "кардиологическим отделением и руководитель Регионального сосудистого центра ЦКБ "
           "(г. Новосибирск). Диагностический блок.")
add_bullet("Кашталап Василий Васильевич — д.м.н., профессор, профессор РАН, заведующий отделом "
           "клинической кардиологии НИИ КПССЗ, профессор кафедры кардиологии и ССХ Кемеровского ГМУ, "
           "помощник вице-президента Российского кардиологического общества. Терапевтический блок.")

add_para("Особенности формата:", bold=True)
add_bullet("Два интерактивных обсуждения между лекторами по сложным вопросам ежедневной практики, "
           "где гайдлайны и клиническая реальность не всегда совпадают.")
add_bullet("4 клинических разбора из реальной практики врача-кардиолога — с анонимным голосованием "
           "аудитории и последующим разбором по актуальным рекомендациям.")
add_bullet("Обновлённые данные по бета-адреноблокаторам у коморбидных пациентов (ХОБЛ — PACE 2026; "
           "ХБП — Wu 2021; СД2 — Maccabi 2025).")

add_para("Очный формат, без онлайн-трансляции. Количество мест ограничено.", italic=True)

doc.add_paragraph()

add_para("Регистрация: до {Дата регистрации} — у Вашего медицинского представителя "
         "{имя представителя, телефон} или по ссылке {ссылка}.",
         bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
